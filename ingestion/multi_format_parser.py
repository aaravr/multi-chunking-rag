"""Multi-Format Document Parser — DOCX, Excel, CSV, JSON, HTML (§10.5).

Extends the parser abstraction layer to support non-PDF document formats.
Each format handler produces ParsedDocument/ParsedPage contracts for
uniform downstream processing.

Optional dependencies (install as needed):
- python-docx: DOCX support
- openpyxl: Excel (.xlsx) support
- beautifulsoup4: HTML support
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
import os
import uuid
from typing import Any, Dict, List, Optional

from agents.contracts import ParsedDocument, ParsedPage
from ingestion.parser_base import BaseParser, register_parser

logger = logging.getLogger(__name__)


# ── DOCX Parser ──────────────────────────────────────────────────────

_DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    _DOCX_AVAILABLE = True
except ImportError:
    pass


class DocxParser(BaseParser):
    """Microsoft Word (.docx) parser."""

    parser_name = "docx"
    supported_formats = ["docx"]

    def parse(self, file_path: str, doc_id: str = "") -> ParsedDocument:
        if not _DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        with open(file_path, "rb") as f:
            data = f.read()
        return self.parse_bytes(data, os.path.basename(file_path), doc_id)

    def parse_bytes(self, data: bytes, filename: str, doc_id: str = "") -> ParsedDocument:
        if not _DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

        if not doc_id:
            sha = hashlib.sha256(data).hexdigest()
            doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sha))

        doc = DocxDocument(io.BytesIO(data))

        # DOCX doesn't have pages — we create logical pages by heading breaks
        current_text: List[str] = []
        current_tables: List[Dict[str, Any]] = []
        current_headings: List[Dict[str, Any]] = []
        pages: List[ParsedPage] = []
        page_num = 1

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                if "Heading" in style_name:
                    # New section → new page
                    if current_text or current_tables:
                        pages.append(ParsedPage(
                            page_number=page_num,
                            text_content="\n".join(current_text),
                            tables=current_tables,
                            headings=current_headings,
                        ))
                        page_num += 1
                        current_text = []
                        current_tables = []
                        current_headings = []

                    level = 1
                    if "2" in style_name:
                        level = 2
                    elif "3" in style_name:
                        level = 3
                    current_headings.append({"text": text, "level": level, "bbox": []})

                current_text.append(text)

            elif tag == "tbl":
                # Table
                for table in doc.tables:
                    if table._element is element:
                        md_lines = []
                        cells = []
                        for i, row in enumerate(table.rows):
                            row_cells = [cell.text.strip() for cell in row.cells]
                            cells.append(row_cells)
                            md_lines.append("| " + " | ".join(row_cells) + " |")
                            if i == 0:
                                md_lines.append("| " + " | ".join("---" for _ in row_cells) + " |")
                        current_tables.append({
                            "markdown": "\n".join(md_lines),
                            "bbox": [],
                            "cells": cells,
                        })
                        break

        # Final page
        if current_text or current_tables:
            pages.append(ParsedPage(
                page_number=page_num,
                text_content="\n".join(current_text),
                tables=current_tables,
                headings=current_headings,
            ))

        if not pages:
            pages.append(ParsedPage(page_number=1, text_content=""))

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            page_count=len(pages),
            pages=pages,
            parser_name=self.parser_name,
            file_format="docx",
        )


# ── Excel Parser ─────────────────────────────────────────────────────

_XLSX_AVAILABLE = False
try:
    from openpyxl import load_workbook
    _XLSX_AVAILABLE = True
except ImportError:
    pass


class ExcelParser(BaseParser):
    """Microsoft Excel (.xlsx) parser — one page per worksheet."""

    parser_name = "xlsx"
    supported_formats = ["xlsx", "xls"]

    def parse(self, file_path: str, doc_id: str = "") -> ParsedDocument:
        if not _XLSX_AVAILABLE:
            raise RuntimeError("openpyxl not installed. Install with: pip install openpyxl")
        with open(file_path, "rb") as f:
            data = f.read()
        return self.parse_bytes(data, os.path.basename(file_path), doc_id)

    def parse_bytes(self, data: bytes, filename: str, doc_id: str = "") -> ParsedDocument:
        if not _XLSX_AVAILABLE:
            raise RuntimeError("openpyxl not installed. Install with: pip install openpyxl")

        if not doc_id:
            sha = hashlib.sha256(data).hexdigest()
            doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sha))

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        pages: List[ParsedPage] = []

        for page_num, sheet_name in enumerate(wb.sheetnames, 1):
            ws = wb[sheet_name]
            rows = []
            text_lines = []

            for row in ws.iter_rows(values_only=True):
                str_row = [str(c) if c is not None else "" for c in row]
                rows.append(str_row)
                text_lines.append("\t".join(str_row))

            # Build markdown table
            md_lines = []
            if rows:
                for i, row in enumerate(rows):
                    md_lines.append("| " + " | ".join(row) + " |")
                    if i == 0:
                        md_lines.append("| " + " | ".join("---" for _ in row) + " |")

            tables = [{
                "markdown": "\n".join(md_lines),
                "bbox": [],
                "cells": rows,
            }] if rows else []

            pages.append(ParsedPage(
                page_number=page_num,
                text_content="\n".join(text_lines),
                tables=tables,
                headings=[{"text": sheet_name, "level": 1, "bbox": []}],
                metadata={"sheet_name": sheet_name},
            ))

        wb.close()

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            page_count=len(pages),
            pages=pages,
            parser_name=self.parser_name,
            file_format="xlsx",
        )


# ── CSV Parser ───────────────────────────────────────────────────────


class CsvParser(BaseParser):
    """CSV file parser — entire file as one page/table."""

    parser_name = "csv"
    supported_formats = ["csv", "tsv"]

    def parse(self, file_path: str, doc_id: str = "") -> ParsedDocument:
        with open(file_path, "rb") as f:
            data = f.read()
        return self.parse_bytes(data, os.path.basename(file_path), doc_id)

    def parse_bytes(self, data: bytes, filename: str, doc_id: str = "") -> ParsedDocument:
        if not doc_id:
            sha = hashlib.sha256(data).hexdigest()
            doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sha))

        text = data.decode("utf-8", errors="replace")
        delimiter = "\t" if filename.lower().endswith(".tsv") else ","
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = [row for row in reader]

        text_lines = []
        md_lines = []
        for i, row in enumerate(rows):
            text_lines.append(delimiter.join(row))
            md_lines.append("| " + " | ".join(row) + " |")
            if i == 0:
                md_lines.append("| " + " | ".join("---" for _ in row) + " |")

        tables = [{
            "markdown": "\n".join(md_lines),
            "bbox": [],
            "cells": rows,
        }] if rows else []

        page = ParsedPage(
            page_number=1,
            text_content="\n".join(text_lines),
            tables=tables,
        )

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            page_count=1,
            pages=[page],
            parser_name=self.parser_name,
            file_format="csv" if delimiter == "," else "tsv",
        )


# ── JSON Parser ──────────────────────────────────────────────────────


class JsonParser(BaseParser):
    """JSON file parser — pretty-prints as text content."""

    parser_name = "json"
    supported_formats = ["json"]

    def parse(self, file_path: str, doc_id: str = "") -> ParsedDocument:
        with open(file_path, "rb") as f:
            data = f.read()
        return self.parse_bytes(data, os.path.basename(file_path), doc_id)

    def parse_bytes(self, data: bytes, filename: str, doc_id: str = "") -> ParsedDocument:
        if not doc_id:
            sha = hashlib.sha256(data).hexdigest()
            doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sha))

        text = data.decode("utf-8", errors="replace")
        try:
            parsed = json.loads(text)
            pretty = json.dumps(parsed, indent=2)
        except json.JSONDecodeError:
            pretty = text

        # If JSON is a list of objects, create a table
        tables = []
        if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
            headers = list(parsed[0].keys())
            rows = [headers]
            for item in parsed:
                rows.append([str(item.get(h, "")) for h in headers])
            md_lines = []
            for i, row in enumerate(rows):
                md_lines.append("| " + " | ".join(row) + " |")
                if i == 0:
                    md_lines.append("| " + " | ".join("---" for _ in row) + " |")
            tables.append({"markdown": "\n".join(md_lines), "bbox": [], "cells": rows})

        page = ParsedPage(
            page_number=1,
            text_content=pretty,
            tables=tables,
        )

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            page_count=1,
            pages=[page],
            parser_name=self.parser_name,
            file_format="json",
        )


# ── HTML Parser ──────────────────────────────────────────────────────

_BS4_AVAILABLE = False
try:
    from bs4 import BeautifulSoup
    _BS4_AVAILABLE = True
except ImportError:
    pass


class HtmlParser(BaseParser):
    """HTML file parser using BeautifulSoup."""

    parser_name = "html"
    supported_formats = ["html", "htm"]

    def parse(self, file_path: str, doc_id: str = "") -> ParsedDocument:
        with open(file_path, "rb") as f:
            data = f.read()
        return self.parse_bytes(data, os.path.basename(file_path), doc_id)

    def parse_bytes(self, data: bytes, filename: str, doc_id: str = "") -> ParsedDocument:
        if not _BS4_AVAILABLE:
            raise RuntimeError("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")

        if not doc_id:
            sha = hashlib.sha256(data).hexdigest()
            doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sha))

        html = data.decode("utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")

        text = soup.get_text(separator="\n", strip=True)

        # Extract headings
        headings = []
        for level in range(1, 7):
            for h in soup.find_all(f"h{level}"):
                headings.append({
                    "text": h.get_text(strip=True),
                    "level": level,
                    "bbox": [],
                })

        # Extract tables
        tables = []
        for table in soup.find_all("table"):
            rows = []
            for tr in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                md_lines = []
                for i, row in enumerate(rows):
                    md_lines.append("| " + " | ".join(row) + " |")
                    if i == 0:
                        md_lines.append("| " + " | ".join("---" for _ in row) + " |")
                tables.append({"markdown": "\n".join(md_lines), "bbox": [], "cells": rows})

        page = ParsedPage(
            page_number=1,
            text_content=text,
            tables=tables,
            headings=headings,
        )

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            page_count=1,
            pages=[page],
            parser_name=self.parser_name,
            file_format="html",
        )


# ── Auto-register parsers ────────────────────────────────────────────

# CSV and JSON always available (stdlib only)
register_parser(CsvParser())
register_parser(JsonParser())

# DOCX requires python-docx
if _DOCX_AVAILABLE:
    register_parser(DocxParser())

# Excel requires openpyxl
if _XLSX_AVAILABLE:
    register_parser(ExcelParser())

# HTML requires beautifulsoup4
if _BS4_AVAILABLE:
    register_parser(HtmlParser())
